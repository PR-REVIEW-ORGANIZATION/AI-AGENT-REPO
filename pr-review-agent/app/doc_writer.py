"""DOCX and JSON output writers for final PR review artifacts."""

from __future__ import annotations

import json
from dataclasses import asdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from app.models import ChangedFile, FileReview, FinalReview, PRMetadata, PrecheckResult


def write_review_docx(
    review: FinalReview,
    metadata: PRMetadata,
    changed_files: list[ChangedFile],
    prechecks: PrecheckResult,
    file_reviews: list[FileReview],
    output_dir: str | Path = "outputs",
) -> Path:
    """Generate the final DOCX artifact with the required section structure."""
    destination_dir = Path(output_dir)
    destination_dir.mkdir(parents=True, exist_ok=True)
    output_path = destination_dir / f"Detailed_PR_Review_PR_{metadata.number}.docx"

    document = Document()

    document.add_heading("1. Document Header", level=1)
    header_table = document.add_table(rows=5, cols=2)
    header_table.style = "Table Grid"
    _set_row(header_table, 0, "Document Name", output_path.name)
    _set_row(header_table, 1, "Repository", metadata.repository)
    _set_row(header_table, 2, "PR Number", str(metadata.number))
    _set_row(header_table, 3, "Generated At (UTC)", datetime.now(timezone.utc).isoformat())
    _set_row(header_table, 4, "Final Decision", review.final_decision)

    document.add_heading("2. Executive Summary", level=1)
    document.add_paragraph(review.executive_summary)

    document.add_heading("3. PR Metadata", level=1)
    metadata_table = document.add_table(rows=10, cols=2)
    metadata_table.style = "Table Grid"
    _set_row(metadata_table, 0, "Title", metadata.title)
    _set_row(metadata_table, 1, "Author", metadata.author)
    _set_row(metadata_table, 2, "State", metadata.state)
    _set_row(metadata_table, 3, "Base Branch", metadata.base_branch)
    _set_row(metadata_table, 4, "Head Branch", metadata.head_branch)
    _set_row(metadata_table, 5, "Commit SHA", metadata.commit_sha)
    _set_row(metadata_table, 6, "Created At", metadata.created_at)
    _set_row(metadata_table, 7, "Updated At", metadata.updated_at)
    _set_row(metadata_table, 8, "Additions / Deletions", f"{metadata.additions} / {metadata.deletions}")
    _set_row(metadata_table, 9, "Changed Files", str(metadata.changed_files_count))

    document.add_heading("4. Business / Functional Context", level=1)
    document.add_paragraph(review.business_functional_context)

    document.add_heading("5. Scope of Review", level=1)
    document.add_paragraph(review.scope_of_review)
    scope_points = [
        f"Deterministic findings: {len(prechecks.issues)}",
        f"Risky files detected: {len(prechecks.risky_files)}",
        f"Large diffs detected: {len(prechecks.large_diff_files)}",
        f"Files with missing patch data: {len(prechecks.missing_patch_files)}",
        f"AI-reviewed files: {sum(1 for item in file_reviews if not item.skipped)}",
    ]
    for point in scope_points:
        document.add_paragraph(point, style="List Bullet")

    document.add_heading("6. Files Changed Summary", level=1)
    files_table = document.add_table(rows=1, cols=6)
    files_table.style = "Table Grid"
    header = files_table.rows[0].cells
    header[0].text = "File"
    header[1].text = "Status"
    header[2].text = "Additions"
    header[3].text = "Deletions"
    header[4].text = "Changes"
    header[5].text = "Flags"

    review_by_file = {review_item.filename: review_item for review_item in file_reviews}
    for changed_file in changed_files:
        row = files_table.add_row().cells
        row[0].text = changed_file.filename
        row[1].text = changed_file.status
        row[2].text = str(changed_file.additions)
        row[3].text = str(changed_file.deletions)
        row[4].text = str(changed_file.changes)
        row[5].text = _build_flags(changed_file, prechecks, review_by_file.get(changed_file.filename))

    document.add_heading("7. What Changed From Existing Behavior", level=1)
    for line in review.behavior_changes:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("8. Final Recommendation", level=1)
    document.add_paragraph(f"Risk Level: {review.risk_level}")
    document.add_paragraph(f"Decision: {review.final_decision}")
    document.add_paragraph(f"Reasoning: {review.reasoning}")

    if review.key_issues:
        issue_table = document.add_table(rows=1, cols=5)
        issue_table.style = "Table Grid"
        issue_header = issue_table.rows[0].cells
        issue_header[0].text = "Severity"
        issue_header[1].text = "Category"
        issue_header[2].text = "Description"
        issue_header[3].text = "Recommendation"
        issue_header[4].text = "Evidence"
        for issue in review.key_issues:
            row = issue_table.add_row().cells
            row[0].text = issue.severity
            row[1].text = issue.category
            row[2].text = issue.description
            row[3].text = issue.recommendation
            row[4].text = issue.evidence or "N/A"

    document.save(output_path)
    return output_path


def write_review_json(
    review: FinalReview,
    metadata: PRMetadata,
    prechecks: PrecheckResult,
    file_reviews: list[FileReview],
    output_dir: str | Path = "outputs",
) -> Path:
    """Write final review payload to JSON artifact."""
    destination_dir = Path(output_dir)
    destination_dir.mkdir(parents=True, exist_ok=True)
    output_path = destination_dir / f"Detailed_PR_Review_PR_{metadata.number}.json"
    payload = {
        "pr_metadata": asdict(metadata),
        "prechecks": prechecks.to_dict(),
        "file_reviews": [review_item.to_dict() for review_item in file_reviews],
        "final_review": review.to_dict(),
    }
    output_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return output_path


def _set_row(table, row_idx: int, key: str, value: str) -> None:
    table.rows[row_idx].cells[0].text = key
    table.rows[row_idx].cells[1].text = value or "N/A"


def _build_flags(
    changed_file: ChangedFile,
    prechecks: PrecheckResult,
    file_review: FileReview | None,
) -> str:
    flags: list[str] = []
    if changed_file.filename in prechecks.risky_files:
        flags.append("risky")
    if changed_file.filename in prechecks.large_diff_files:
        flags.append("large-diff")
    if changed_file.filename in prechecks.missing_patch_files:
        flags.append("missing-patch")
    if file_review and file_review.skipped:
        flags.append("ai-skipped")
    return ", ".join(flags) if flags else "none"
